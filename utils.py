import os
import re
import time

import streamlit as st

from openai import OpenAI
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from dotenv import load_dotenv


# Load OpenAI API Key
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


class DocumentReader:
    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.document = None

    def read_docx(self):
        try:
            self.document = Document(self.file_path)
            full_text = []
            for para in self.document.paragraphs:
                full_text.append(para.text)
            return "\n".join(full_text)
        except Exception as e:
            print(e)
            return None


class DocumentProcessor:
    def __init__(self, document: str) -> None:
        self.document = document

    def chunk_document(self, chunk_size: int) -> list:
        chunks = []
        text = self.document
        while text:
            boundary = min(len(text), chunk_size)
            # Search for the last sentence end near the boundary
            if boundary < len(text):
                match = list(re.finditer(r"\.\n|[\.!?] [A-Z]", text[:boundary + 50]))
                if match:
                    boundary = match[-1].start() + 1
            chunks.append(text[:boundary].strip())
            text = text[boundary:]
        return chunks


class Proofreader:
    def __init__(
        self,
        document_chunks: list,
        model: str,
        temperature: float,
        max_tokens: int = 2000,
        n: int = 1,
        stop=None,
    ) -> None:
        self.document_chunks = document_chunks
        self.model = model
        self.temperature = temperature
        self.max_tokens = max_tokens
        self.n = n
        self.stop = stop
        self.outputs = {}
        self.output_text = ""
        self.leads = []
        self.quotes = []
        self.titles = []
        self.tags_from_list = []
        self.tags = []

    def get_openai_response(self, system_message: str, user_message: str) -> str:
        response = client.chat.completions.create(
            model=self.model,
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_message},
                ],
            temperature=self.temperature,
            n=self.n,
            stop=self.stop,
        )
        return response.choices[0].message.content.strip()

    def proofread(self, chunk: str, system_message: str) -> dict:
        # TODO: Create config dict with prompt templates
        """
        Proofreads a chunk of text.
        """
        st.write("Poprawiam fragment...")
        user_message = f'''\
        Poniższy tekst jest fragmentem artykułu, który ma zostać opublikowany w gazecie, lub całym artykułem.
        Przeczytaj poniższy tekst i dokonaj korekty błędów interpunkcyjnych, ortograficznych,
        gramatycznych i składniowych oraz stylistycznych.
        Nie streszczaj.
        Zachowaj całą treść tekstu, tylko popraw go językowo.
        Nie ingeruj w strukturę akapitów.
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi podaj tylko sam poprawiony tekst.
        '''

        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text

    def create_heading(self, chunk: str, system_message: str) -> dict:
        # TODO: Create config dict with prompt templates
        """
        Creates heading a chunk of text.
        """
        st.write("Tworzę nagłówek...")
        user_message = f'''
        Przeczytaj poniższy fragment tekstu i wymyśl do niego propozycję ciekawego nagłówka.
        WAŻNE: długość nagłówka nie może przekraczać 4 słów.
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi podaj tylko sam nagłówek.
        '''

        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text

    def get_quotes(self, chunk: str, system_message: str) -> str:
        """
        Gets quotes from a chunk of text.
        """
        st.write("Wyciągam wyimy...")
        user_message = f'''
        Przeczytaj poniższy tekst i wybierz z niego 5 cytatów, które mogą być najbardziej interesujące dla czytelników.
        Muszą to być dosłowne cytaty z tekstu i obejmować co najmniej jedno zdanie.
        Ważne: Długość cytatu NIE MOŻE przekraczać 250 znaków!!!
        Tekst do analizy:"""{chunk}"""
        W odpowiedzi wypisz tylko same wybrane cytaty.
        Proponowane cytaty:
        '''
        # TODO: more creativity?
        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text.split("\n")

    def create_titles(self, system_message: str) -> str:
        """
        Creates propositions of a title for a text based on summary.
        """
        st.write("Tworzę tytuły...")
        user_message = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższy tekst i napisz trzy propozycje interesującego i przyciągającego uwagę tytułu.
        Ważne: Długość tytułu NIE MOŻE przekraczać 150 znaków!!!
        Przykładowe, dobre tytuły:
        - "Biden rezygnuje. Amerykańska kampania wyborcza właśnie się zresetowała"
        - "Obecny prezydent USA poświęcił osobistą ambicję dla dobra amerykańskiej demokracji.
        I dla dobra demokracji na świecie"
        - "Po pandemii widać turystykę i wojnę - nowe dane o rezerwacjach na platformach internetowych"
        - "Emmanuel Macron podał piłkę Marine Le Pen"
        - "Legalizacja związków partnerskich leży w interesie koalicji rządzącej, w tym także PSL"

        Tekst do analizy:"""{" ".join(self.document_chunks)}"""

        Proponowane tytuły:
        '''
        # TODO: more creativity?
        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text.split("\n")

    def create_leads(self, system_message: str) -> str:
        """
        Creates propositions of lead for a text based on summary.
        """
        st.write("Tworzę leady...")
        # TODO: temporary solution. We have to deal with the summary length, but how?
        user_message = f'''
        Jesteś doświadczonym redaktorem.
        Przeczytaj poniższe streszczenie i napisz trzy propozycje interesującego i przyciągającego uwagę leadu.
        Lead może np. zawierać mocne tezy z tekstu.
        Ważne: Długość leadu NIE MOŻE przekraczać 250 znaków!!!
        Przykładowe, dobre leady:
        - "Negocjując z Chinami, podbijamy stawkę Waszyngtonowi oraz zmuszamy Berlin, by bardziej uwzględniał nas
        w swoim równaniu strategicznym"
        - "Zamiast likwidować IPN, wyłączmy z niego i przenieśmy do prokuratury powszechnej pion śledczy,
        a to, co zostanie, wcielmy do Polskiej Akademii Nauk na prawach jednego z jej instytutów"
        "Europie grozi deficyt demokracji. Potężnym problemem jest dominacja państw Europy Zachodniej i
        Północnej nad peryferiami z Europy Południowej i Środkowej. Jeśli UE chce przetrwać,
        powinno się ograniczyć dominację państw najsilniejszych"

        Tekst do analizy:"""{" ".join(self.document_chunks)}"""

        W odpowiedzi wypisz tylko same wybrane leady.
        Proponowane leady:
        '''

        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text.split("\n")

    def create_tags_from_list(self, tag_list: str, system_message: str) -> list:
        """
        Select tags from a tag list for a text based on summary.
        """
        st.write("Tworzę tagi z listy...")
        # TODO: temporary solution. We have to deal with the summary length, but how?
        prompt = f'''
        Przeczytaj poniższy tekst i napisz do niego maksymalnie trzy propozycje
        najbardziej pasujących tagów wybranych z następującej listy: {tag_list}
        Tekst do analizy:"""{" ".join(self.document_chunks)}"""
        W odpowiedzi wypisz tylko same wybrane tagi.
        Wybrane tagi:

        '''

        response_text = self.get_openai_response(
            system_message=system_message, user_message=prompt
        )
        return response_text.split("\n")

    def create_tags(self, tag_list, system_message: str) -> list:
        """
        Creates tags for a text based on summary.
        """
        st.write("Tworzę tagi swobodne...")
        # TODO: temporary solution. We have to deal with the summary length, but how?
        user_message = f'''
        Przeczytaj poniższy tekst i napisz dla niego pięć propozycji tagów.
        Format tagu: "#wybory w USA".
        Tag musi mieć 2-4 słowa.
        Wśród tagów nie może być tagi z listy: {tag_list}.
        Tekst do analizy:"""{" ".join(self.document_chunks)}"""
        W odpowiedzi wypisz tylko same wybrane tagi.
        Wybrane tagi:
        '''

        response_text = self.get_openai_response(
            system_message=system_message, user_message=user_message
        )
        return response_text.split("\n")

    def process_document(self, system_message: str, tag_list: list, midtitles: bool = False) -> None:
        st.write("Zaczynam redagować dokument.")
        # Start counting time of method execution
        time_start = time.time()
        for chunk in self.document_chunks:
            st.write("Redaguję fragment zaczynający się od:", chunk[:10])
            if midtitles:
                subtitle = "\n\n" + self.create_heading(chunk, system_message) + "\n\n"
                chunk = subtitle + self.proofread(chunk, system_message)
            else:
                chunk = self.proofread(chunk, system_message)
            self.output_text += chunk
            st.write("Czas od rozpoczęcia procesu:", time.time() - time_start)
            st.write("Czas od rozpoczęcia procesu:", time.time() - time_start)
            for quote in self.get_quotes(chunk, system_message):
                self.quotes.append(quote)
            st.write("Czas od rozpoczęcia procesu:", time.time() - time_start)
        self.titles = self.create_titles(system_message)
        self.leads = self.create_leads(system_message)
        self.tags_from_list = self.create_tags_from_list(tag_list, system_message)
        self.tags = self.create_tags(tag_list, system_message)
        self.outputs = {
            "titles": self.titles,
            "leads": self.leads,
            "tags_from_list": self.tags_from_list,
            "tags": self.tags,
            "quotes": self.quotes,
            "output_text": self.output_text,
        }


class DocumentStyler:
    def __init__(self) -> None:
        pass

    def set_style(self, document: Document) -> None:
        normal_style = document.styles["Normal"]
        font = normal_style.font
        font.name = "Calibri"
        font.size = Pt(11)
        paragraph_format = normal_style.paragraph_format
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph_format.line_spacing = 1.0
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(6)

        title_style = document.styles.add_style("Custom Heading", WD_STYLE_TYPE.PARAGRAPH)
        font = title_style.font
        font.name = "Calibri"
        font.size = Pt(11)
        font.bold = True
        paragraph_format = title_style.paragraph_format
        paragraph_format.space_before = Pt(6)
        paragraph_format.space_after = Pt(6)

        return document


class DocumentWriter:
    """
    Writes a list of texts to a docx file.
    """

    def __init__(self, file_path: str) -> None:
        self.file_path = file_path
        self.document = None

    def write_document(self, output: dict) -> None:
        self.document = Document()
        styler = DocumentStyler()
        styler.set_style(self.document)

        # Add titles section
        self.document.add_paragraph("Tytuły", style="Custom Heading")
        for title in output["titles"]:
            title = title.replace('"', "")
            self.document.add_paragraph(title)

        self.document.add_paragraph("\n\n")

        # Add leads section
        self.document.add_paragraph("Leady", style="Custom Heading")
        for lead in output["leads"]:
            lead = lead.replace('"', "")
            self.document.add_paragraph(lead)

        self.document.add_paragraph("\n\n")

        # Add tags section
        self.document.add_paragraph("Tagi", style="Custom Heading")
        # TODO: Temporary cleaning data solution
        list_tags = ", ".join(output["tags_from_list"]).replace(",, ", ", ")
        list_tags = "Tagi: " + list_tags.replace(", , ", ", ")
        self.document.add_paragraph(list_tags)
        tags = ", ".join(output["tags"]).replace(",, ", ", ")
        tags = "# " + tags.replace(", , ", ", ")
        self.document.add_paragraph(tags)

        self.document.add_paragraph("\n\n")

        # Add quotes section
        self.document.add_paragraph("Cytaty", style="Custom Heading")
        for quote in output["quotes"]:
            quote = quote.replace('"', "")
            self.document.add_paragraph(quote)

        self.document.add_paragraph("\n\n")

        # Add text section
        self.document.add_paragraph("Poprawiony tekst", style="Custom Heading")
        self.document.add_paragraph(output["output_text"])

        self.document.save(self.file_path)
